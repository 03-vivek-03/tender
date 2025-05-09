import streamlit as st
import cohere
import PyPDF2
from docx import Document
from io import BytesIO
from docx.shared import Pt
import gspread
from google.oauth2.service_account import Credentials
from datetime import datetime
import pytz
import re
# Initialize Cohere client
co = cohere.ClientV2(api_key="okYrKAw1OPZoMnOSCR6rUVO2cbSulB4gCmuo04UY")  # Replace with your key

def log_to_google_sheet(filename, file_data, extracted_text):
    scopes = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
    credentials = Credentials.from_service_account_info(
        st.secrets["google_sheets"],
        scopes=scopes
    )
    client = gspread.authorize(credentials)
    sheet = client.open("TenderUsageLogs").sheet1

    # Add headers if sheet is empty
    if sheet.row_count == 0 or sheet.cell(1, 1).value != "Timestamp":
        sheet.insert_row(["Timestamp", "Filename", "File Size (KB)", "Text Length", "User IP"], 1)

    # Current time in IST
    ist = pytz.timezone("Asia/Kolkata")
    timestamp = datetime.now(ist).strftime("%Y-%m-%d %H:%M:%S")

    # User IP (fallback to N/A)
    user_ip = st.request.remote_addr if hasattr(st, 'request') else "N/A"

    # File size in KB
    file_data.seek(0, 2)
    file_size_kb = round(file_data.tell() / 1024, 2)
    file_data.seek(0)

    # Extracted text length
    text_length = len(extracted_text.strip())

    # Append the log
    sheet.append_row([timestamp, filename, f"{file_size_kb} KB", text_length, user_ip])
    return True

def extract_text_from_pdf(pdf_file):
    reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def generate_table_word(summary_text):
    # Split lines and extract the first heading (any level)
    lines = summary_text.splitlines()
    heading = next((l.strip().lstrip('#').strip() for l in lines if l.strip().startswith('#')), 'Table')

    # Parse keys and their bullet values
    data = []
    key = None
    values = []
    for line in lines:
        stripped = line.lstrip('#').strip()
        if re.match(r'^\*\*.*\*\*$', stripped):  # **Key**
            if key:
                data.append((key, values))
            key = stripped.strip('*').strip()
            values = []
        elif stripped.startswith('-') and key:  # - bullet under key
            values.append(stripped.lstrip('-').strip())
    if key:
        data.append((key, values))

    # Create Word document and add heading
    doc = Document()
    title_para = doc.add_heading(level=1)
    run_title = title_para.add_run(heading)
    run_title.bold = True
    run_title.font.size = Pt(16)

    doc.add_paragraph()  # spacing

    # Build table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Description'
    for cell in hdr_cells:
        for p in cell.paragraphs:
            p.alignment = 1  # Center
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)

    # Populate table
    for key, vals in data:
        row_cells = table.add_row().cells
        cell_key = row_cells[0]
        p_key = cell_key.paragraphs[0]
        run_key = p_key.add_run(key)
        run_key.bold = True
        run_key.font.size = Pt(11)

        cell_val = row_cells[1]
        if vals:
            for v in vals:
                p = cell_val.add_paragraph(style='List Bullet')
                parts = re.split(r'(\*\*[^*]+\*\*)', v)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part.strip('*'))
                        run.bold = True
                    else:
                        p.add_run(part)
        else:
            cell_val.text = ''

    # Save to BytesIO
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def stream_summary_from_cohere(text):
    prompt = (
        """You are an expert in analyzing and summarizing government and institutional tender documents.

            Summarize the following tender document by extracting and presenting all important and relevant information that may be present. Only include the sections that are explicitly mentioned or applicable in the document. **Do not include sections that are not present, not mentioned, or not relevant to the specific tender type.**

            Extract details under the following categories **only if available**:
            - Tender Name
            - Tender Reference Number and ID  
            - Name of the Issuing Organization or Authority  
            - Tender Fee (amount, mode of payment)  
            - EMD (Earnest Money Deposit) Details (amount, mode of payment)  
            - Estimated Tender Value or Project Cost  
            - Pre-bid Meeting Dates, Venue, and Registration/Link  
            - Tender Meeting Dates and Venues (if different from Pre-bid)  
            - Scope of Work  
            - Modules or Work Packages  
            - Workforce Requirements (specify onsite manpower and training manpower, if any)  
            - Human Resource Details  
            - Technical and Financial Eligibility Criteria  
            - Technical and Financial Marking/Scoring Criteria  
            - Performance Security Requirements  
            - Implementation Timeline and Phases (Turnaround Time or TAT)  
            - Contract Duration/Period  
            - Project Location(s)  
            - Existing IHMS or Software Application Details (if mentioned)  
            - Payment Terms and Schedule  
            - Submission Method (Online, Physical, or Hybrid)  
            - Selection Methodology (e.g., QCBS, L1)  
            - Cloud Service Provider (CSP) Details (if applicable)  
            - Hardware Details (especially for hospital/lab tenders—CT/MRI/X-ray/Pathology equipment)  
            - Technical Specifications  
            - Radiology/Pathology Scope (if applicable)  
            - Checklists (All the documents required, if provided)  
            - Declarations, Undertakings, and Affidavits  
            - OEM (Original Equipment Manufacturer) Document Requirements  
            - Penalty Clauses and Bidder Obligations  
            - Financial Bid Structure  
            - Viability Gap Funding (VGF)  
            - Special Purpose Vehicle (SPV) clauses  
            - Land Border Sharing Clause  
            - Mode of Payments for Tender Fee, EMD, and Other Charges  
            - Contact Details of the Tender Issuer (email, phone, address)

            Present the summary in a clean, organized format using clear headings or bullet points. Again, include **only the sections that are actually present in the document** and dont say not mentioned in the document, instead skip that section.
            
            At last give me these details seperate again: Tender Name, Tender Type (HIMS, Radiology Lab etc.), Tender registration start date and end date)
            
            Tender Document:\n\n"""
        f"{text}"
    )

    response = co.chat_stream(
        model="command-a-03-2025",
        messages=[{"role": "user", "content": prompt}]
    )

    for chunk in response:
        if chunk and chunk.type == "content-delta":
            yield chunk.delta.message.content.text

# Set page config
st.set_page_config(page_title="Tender Summarizer", page_icon="📄")

# UI content
st.title("📄 Tender Document Summarizer")
st.markdown("Upload a **tender PDF** and get a concise summary with all key information extracted.")

uploaded_file = st.file_uploader("Upload PDF", type=["pdf"])

if uploaded_file is not None:
    with st.spinner("File uploaded successfully!✅"):
        text = extract_text_from_pdf(uploaded_file)
        log_to_google_sheet(uploaded_file.name, uploaded_file, text)

    if len(text.strip()) < 100:
        st.error("The uploaded PDF has very little text or is not extractable.")
    else:
        st.success("Generating summary...\n")

        # Add a placeholder to dynamically update the summary
        summary_placeholder = st.empty()
        if "summary" not in st.session_state:
            summary_text = ""
            for chunk in stream_summary_from_cohere(text):
                summary_text += chunk
                summary_placeholder.markdown(summary_text)
            st.session_state["summary"] = summary_text
        else:
            summary_text = st.session_state["summary"]
            summary_placeholder.markdown(summary_text)

        # # Generate Word document with formatting
        # doc = Document()
        # style = doc.styles["Normal"]
        # style.font.size = Pt(11)

        # doc.add_heading("Tender Summary", level=1)

        # # Process the summary line by line
        # for line in summary_text.splitlines():
        #     line = line.strip()
        #     if not line:
        #         doc.add_paragraph("")  # Preserve blank lines
        #         continue

        #     # Handle markdown-style headings
        #     if line.startswith("####"):
        #         doc.add_heading(line.replace("####", "").strip(), level=4)
        #     elif line.startswith("###"):
        #         doc.add_heading(line.replace("###", "").strip(), level=3)
        #     else:
        #         # Split the line based on '**' to handle bold text
        #         paragraph = doc.add_paragraph()
        #         segments = line.split('**')

        #         for i, segment in enumerate(segments):
        #             if i % 2 == 1:  # This part should be bold (because it's between '**')
        #                 paragraph.add_run(segment).bold = True
        #             else:  # Normal text (no bold)
        #                 paragraph.add_run(segment)

        # # Save to BytesIO buffer
        # word_buffer = BytesIO()
        # doc.save(word_buffer)

        # word_buffer.seek(0)

        # # Download button
        # st.download_button(
        #     label="📅 Download Summary",
        #     data=word_buffer,
        #     file_name=f"{uploaded_file.name.rsplit('.', 1)[0]}_Summary.docx", # Prefix added here
        #     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        # )


        # Generate table format Word document
        table_buffer = generate_table_word(summary_text)

        # Second download button for table format
        st.download_button(
            label="📊 Download Summary in Table format",
            data=table_buffer,
            file_name=f"{uploaded_file.name.rsplit('.', 1)[0]}_Table_Summary.docx",  # Prefix from uploaded filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.info("Please upload a tender PDF file to begin.")

# Footer and close centered div
st.markdown("---")
st.markdown(
    "<p style='text-align:center; color: gray;'>Designed by Medimaze AI Team</p></div>",
    unsafe_allow_html=True,
)
